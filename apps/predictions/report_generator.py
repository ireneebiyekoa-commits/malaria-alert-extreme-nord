"""
Génération de rapports Word (.docx) — page Prévisions.

Charge TEMPLATE.docx à la racine du projet (entête institutionnel + logos
préremplis) et insère après celui-ci un saut de page suivi du contenu
dynamique (synthèse, tableaux, analyse IA).
"""
import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# matplotlib est OPTIONNEL : si absent (déploiement allégé PythonAnywhere/Render),
# le rapport est généré sans graphique mais reste pleinement fonctionnel.
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    _MATPLOTLIB_OK = True
except ImportError:
    _MATPLOTLIB_OK = False

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

TEMPLATE_PATH = settings.BASE_DIR / 'TEMPLATE.docx'


def _load_template() -> Document:
    """Charge le template Word s'il existe, sinon un document vierge."""
    if TEMPLATE_PATH.exists():
        try:
            doc = Document(str(TEMPLATE_PATH))
            logger.info(f"Template chargé : {TEMPLATE_PATH}")
            return doc
        except Exception as exc:
            logger.warning(f"Template illisible ({exc}), document vierge utilisé.")
    else:
        logger.warning(f"TEMPLATE.docx introuvable à {TEMPLATE_PATH}, document vierge.")
    return Document()


def _add_page_break(doc: Document):
    """Insère un saut de page propre après le contenu existant."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _safe_style(table, style_name: str, fallback: str = 'Table Grid'):
    """Applique un style à un tableau, avec fallback si le style n'existe pas."""
    try:
        table.style = style_name
    except KeyError:
        try:
            table.style = fallback
        except KeyError:
            pass


def generer_rapport_previsions(
    district: str,
    algorithme: str,
    horizon: int,
    previsions: List[Dict[str, Any]],
    historique: List[Dict[str, Any]],
    metriques: Dict[str, float],
    analyse_ia: str,
    auteur: str = "Équipe Suivi-Évaluation du GTR",
) -> io.BytesIO:
    """Génère un rapport Word complet et retourne un buffer BytesIO."""
    doc = _load_template()

    # Si pas de template, fixer les marges
    if not TEMPLATE_PATH.exists():
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ============================================================
    # Saut de page entre le contenu institutionnel du template
    # et le contenu dynamique du rapport
    # ============================================================
    _add_page_break(doc)

    # ============================================================
    # En-tête du rapport (titre + métadonnées)
    # ============================================================
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RAPPORT DE PRÉVISION ÉPIDÉMIQUE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0E, 0x47, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(
        f"District de {district.replace('District ', '')}  ·  "
        f"Algorithme {algorithme}  ·  Horizon {horizon} mois"
    )
    s.italic = True
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(0x5A, 0x68, 0x78)

    doc.add_paragraph()

    # Bloc métadonnées (encadré)
    meta_table = doc.add_table(rows=2, cols=2)
    _safe_style(meta_table, 'Light Shading Accent 1', 'Table Grid')
    meta_table.rows[0].cells[0].text = "Date d'émission"
    meta_table.rows[0].cells[1].text = datetime.now().strftime("%d/%m/%Y à %H:%M")
    meta_table.rows[1].cells[0].text = "Auteur"
    meta_table.rows[1].cells[1].text = auteur

    # Mise en gras de la colonne label
    for row in meta_table.rows:
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # ============================================================
    # Section 1 — Synthèse
    # ============================================================
    doc.add_heading("1. Synthèse des prévisions", level=1)
    doc.add_paragraph(
        f"Ce rapport présente les prévisions d'incidence palustre générées par le modèle "
        f"{algorithme} pour le {district} aux horizons 1, 2 et 3 mois. Les prévisions sont "
        f"comparées aux seuils d'alerte épidémique (P75 et P90) définis selon les "
        f"recommandations de l'Organisation Mondiale de la Santé (OMS, 2014). "
        f"Les niveaux d'alerte sont codés VERT (situation normale), ORANGE (élevé) et "
        f"ROUGE (critique)."
    )

    # ============================================================
    # Section 2 — Tableau des prévisions
    # ============================================================
    doc.add_heading("2. Prévisions détaillées", level=1)

    table = doc.add_table(rows=1, cols=5)
    _safe_style(table, 'Light Grid Accent 1', 'Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Horizon"
    hdr[1].text = "Date cible"
    hdr[2].text = "Incidence prédite (/1000)"
    hdr[3].text = "Cas attendus"
    hdr[4].text = "Niveau d'alerte"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for p in previsions:
        row = table.add_row().cells
        row[0].text = f"h = {p['horizon']}"
        row[1].text = str(p.get('date', '—'))
        row[2].text = f"{p['incidence']:.2f}"
        row[3].text = f"{int(p.get('cas', 0)):,}".replace(',', ' ')
        niv = p.get('niveau', 'vert')
        label_niv = {
            'vert': 'NORMAL',
            'orange': 'ÉLEVÉ',
            'rouge': 'CRITIQUE',
        }.get(niv, str(niv).upper())
        row[4].text = label_niv

    doc.add_paragraph()

    # ============================================================
    # Section 3 — Graphique (si matplotlib disponible)
    # ============================================================
    if _MATPLOTLIB_OK:
        doc.add_heading("3. Visualisation graphique", level=1)
        img_buffer = _generer_graphique(historique, previsions, district, algorithme)
        if img_buffer:
            doc.add_picture(img_buffer, width=Inches(6.0))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(
                f"Figure 1 — Évolution observée (12 derniers mois) et prévisions {algorithme} "
                f"pour le {district}."
            )
            r.italic = True
            r.font.size = Pt(9)
        doc.add_paragraph()
    # Si pas de matplotlib, on passe directement à la section 4 (renumérotée en 3)

    section_perf = "4" if _MATPLOTLIB_OK else "3"
    section_analyse = "5" if _MATPLOTLIB_OK else "4"

    # ============================================================
    # Section 4 — Performances
    # ============================================================
    doc.add_heading(f"{section_perf}. Performance du modèle", level=1)

    perf_table = doc.add_table(rows=1, cols=3)
    _safe_style(perf_table, 'Light Grid Accent 1', 'Table Grid')
    h = perf_table.rows[0].cells
    h[0].text = "RMSE"
    h[1].text = "MAE"
    h[2].text = "R²"
    for c in h:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    row = perf_table.add_row().cells
    row[0].text = f"{metriques.get('rmse', 0):.3f}"
    row[1].text = f"{metriques.get('mae', 0):.3f}"
    row[2].text = f"{metriques.get('r2', 0):.3f}"

    doc.add_paragraph(
        "Métriques officielles validées en walk-forward sur 5 plis annuels (2021-2025), "
        "agrégées sur l'ensemble des 32 districts sanitaires de la région."
    )

    doc.add_paragraph()

    # ============================================================
    # Section 5 — Analyse interprétative
    # ============================================================
    doc.add_heading(f"{section_analyse}. Analyse interprétative", level=1)

    analyse = (analyse_ia or '').strip()
    if not analyse:
        doc.add_paragraph(
            "Aucune analyse interprétative n'a été générée. Pour obtenir une analyse "
            "automatique, cliquez sur le bouton « Générer l'analyse » dans la page "
            "Prévisions avant de télécharger ce rapport."
        ).runs[0].italic = True
    else:
        for line in analyse.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                p = doc.add_paragraph()
                p.add_run(line.strip('*')).bold = True
            else:
                # Gestion basique du gras inline **xxx**
                p = doc.add_paragraph()
                _add_inline_bold(p, line)

    # ============================================================
    # Pied de rapport
    # ============================================================
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = foot.add_run(
        "_______\n"
        "Rapport généré automatiquement par la plateforme d'alerte précoce du paludisme. "
        "Document confidentiel à usage interne."
    )
    f.italic = True
    f.font.size = Pt(9)
    f.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_inline_bold(paragraph, text: str):
    """Ajoute un texte à un paragraphe en gérant le gras Markdown **inline**."""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _generer_graphique(historique, previsions, district, algorithme):
    """Génère un graphique matplotlib des prévisions et le retourne en BytesIO."""
    if not _MATPLOTLIB_OK:
        return None
    try:
        dates_hist = [h['date'] for h in historique[-12:]]
        inc_hist = [h['incidence'] for h in historique[-12:]]

        dates_pred = [p['date'] for p in previsions]
        inc_pred = [p['incidence'] for p in previsions]

        fig, ax = plt.subplots(figsize=(9, 4))
        ax.plot(dates_hist, inc_hist, marker='o', color='#0e476e',
                linewidth=2, label='Observé', markersize=6)
        ax.plot(dates_pred, inc_pred, marker='s', color='#dc3545',
                linewidth=2, linestyle='--', label=f'Prévision {algorithme}', markersize=8)

        ax.set_xlabel('Mois')
        ax.set_ylabel('Incidence (/1000 hab.)')
        ax.set_title(f"{district} — Prévisions {algorithme}",
                     fontsize=12, fontweight='bold')
        ax.legend(loc='upper left')
        ax.grid(alpha=0.3)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=130, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning(f"Échec génération graphique : {exc}")
        return None
